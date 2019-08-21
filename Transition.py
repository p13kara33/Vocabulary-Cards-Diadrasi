import xlrd
from docx import Document

#document = Document('Cards.docx')
document = Document()

print('Για να τυπωθει το αρχείο σας ονομάστε το "Lexi.xls" και προσθέστε το στον φάκελο που έχετε ανοίξει')
print(' ')
print(' ')

print('Αναμένετε μέχρι να ολοκληρωθεί το πρόγραμμα, και ανοίξτε το αρχείο "Cards1.docx"')
print(' ')
print(' ')


workbook = xlrd.open_workbook('Lexi.xls')
sheet = workbook.sheet_by_name('Lexico001')


for row in range(sheet.nrows):
        for column in range(sheet.ncols):
            count = row
            
#print(count)
#print (count)
a = count//12 #to synolo ton liston mas
#print (a ,'a')
up = count%12

peris = count - (a * 12)+1 
#print ('p=', peris)



table1 = document.add_table(rows = count+((peris+5)*2), cols = 2)

print(count+(peris*2))
def height_rule(self, value):
        self._tr.trHeight_hRule = 3,5

c=0#metritis 
e=0
st='text'
#print (st, type(st))

#--giving the ability to print Definition Translation synonms ..etc behind each word
print('Τι θέλετε να τυπωθεί πίσω από τις λέξεις, για Definition πατήστε 1, για Translation πατήστε 2, για Synonyms 3, για Derivatives 4 και για Example 5')
x = input()
x = int(x)

#--wrong case 
while x<0 or x>5:
        print('Λαθός πληκτρολόγηση :(')
        print('Τι θέλετε να τυπωθεί πίσω από τις λέξεις, για Definition πατήστε 1, για Translation πατήστε 2, για Synonyms 3, για Derivatives 4 και για Example 5')
        x = input()   
        

while e<(a*12):
    #print(e)
    for i in range (0, 6, 1):
        rows = c+i
        #rowd = 
        w = sheet.cell(rows, 0)
        w= str(w)
        w = w[6:-1]
        #print(w)
        d = sheet.cell(rows, x)
        d=str(d)
        d = d[6:-1]
        #print(d)
        cell0 = table1.cell(rows, 0)
        cell0.text = w 
        cell1 = table1.cell(rows+6, 1)
        cell1.text = d
        e += 1
            
    for j in range (6, 12, 1):
        rows = c+j
        w = sheet.cell(rows, 0)
        w= str(w)
        w = w[6:-1]
        #print (w)
        d = sheet.cell(rows, x)
        d=str(d)
        d = d[6:-1]
        #print (d)
        cell0 = table1.cell(rows-6, 1)
        cell0.text = w 
        cell1 = table1.cell(rows, 0)
        cell1.text = d       
        e +=1
        print('=== :D ===')   
    c +=12
r=0
for l in range(0, peris, 1):
        rows = c + l    
        w = sheet.cell(rows, 0)
        w = str (w)
        w = w[6:-1]
        d = sheet.cell(rows, x)
        d=str(d)
        d = d[6:-1]
        #print(d)
        cell0 = table1.cell(rows, 0)
        cell0.text = w
        cell1 = table1.cell(rows+6, 1)
        cell1.text = d
             
document.save('Cards01.docx')
